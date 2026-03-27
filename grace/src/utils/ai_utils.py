from pathlib import Path
from typing import Any, Dict, List, Union

from src.tools.filemanager.filemanager import FileManager
from src.types.config import AIConfig

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pandas as pd
except ImportError:
    pd = None


def estimate_tokens(text: str) -> int:
    # Rough estimation: 1 token ≈ 4 characters for English text
    return len(text) // 4


def extract_pdf_content(file_path: Path) -> str:
    if PdfReader is None:
        raise ImportError(
            "PyPDF2 is required to read PDF files. Install it with: uv add PyPDF2"
        )

    try:
        with open(file_path, "rb") as file:
            pdf_reader = PdfReader(file)
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"### Page {page_num}\n{text}")
            return "\n\n".join(text_content)
    except Exception as e:
        raise Exception(f"Failed to extract PDF content: {str(e)}")


def extract_docx_content(file_path: Path) -> str:
    if Document is None:
        raise ImportError(
            "python-docx is required to read DOCX files. Install it with: uv add python-docx"
        )

    try:
        doc = Document(file_path)
        text_content = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_content.append("\n" + "\n".join(table_text) + "\n")

        return "\n\n".join(text_content)
    except Exception as e:
        raise Exception(f"Failed to extract DOCX content: {str(e)}")


def extract_excel_content(file_path: Path) -> str:
    if pd is None:
        raise ImportError(
            "pandas is required to read Excel files. Install it with: uv add pandas openpyxl"
        )

    try:
        excel_file = pd.ExcelFile(file_path)
        text_content = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            text_content.append(f"### Sheet: {sheet_name}\n")

            if not df.empty:
                df = df.fillna("")
                table_str = df.to_string(index=False)
                text_content.append(table_str)
            else:
                text_content.append("(Empty sheet)")

        return "\n\n".join(text_content)
    except Exception as e:
        raise Exception(f"Failed to extract Excel content: {str(e)}")


def combine_markdown_files(
    filemanager: FileManager, markdown_files: List[Path], sendAsString: bool = False
) -> Union[str, List[str]]:
    combined_content: List[str] = []
    for file_path in markdown_files:
        try:
            file_extension = file_path.suffix.lower()
            if file_extension == ".pdf":
                content = extract_pdf_content(file_path)
            elif file_extension in [".docx", ".doc"]:
                content = extract_docx_content(file_path)
            elif file_extension in [".xlsx", ".xls"]:
                content = extract_excel_content(file_path)
            else:
                content = filemanager.read_file(file_path)

            combined_content.append(
                f"## Content from {file_path.name}\n\n{content}\n\n"
            )
        except Exception as e:
            combined_content.append(
                f"## Error reading {file_path.name}\n\nError: {str(e)}\n\n"
            )
    if sendAsString:
        return "\n".join(combined_content)
    return combined_content


def estimate_token_usage(
    filemanager: FileManager, markdown_files: List[Path], config: AIConfig
) -> Dict[str, Any]:
    try:
        combined_content = combine_markdown_files(
            filemanager, markdown_files, sendAsString=True
        )
        estimated_input_tokens = estimate_tokens(combined_content)
        estimated_total_tokens = estimated_input_tokens + config.max_tokens
        return {
            "estimated_input_tokens": estimated_input_tokens,
            "max_output_tokens": config.max_tokens,
            "estimated_total_tokens": estimated_total_tokens,
            "model": config.model_id,
        }
    except Exception as e:
        return {"error": str(e)}


def chunk_content_by_tokens(
    content_list: List[Dict[str, str]], max_tokens_per_chunk: int = 80000
) -> List[List[Dict[str, str]]]:
    """
    Chunk content into smaller batches based on estimated token count.

    Args:
        content_list: List of dicts with 'url' and 'content' keys
        max_tokens_per_chunk: Maximum tokens per chunk (default 80k)

    Returns:
        List of content chunks, where each chunk is a list of content dicts
    """
    chunks = []
    current_chunk = []
    current_tokens = 0

    for item in content_list:
        item_tokens = estimate_tokens(item["content"])

        # If adding this item would exceed the limit, start a new chunk
        if current_tokens + item_tokens > max_tokens_per_chunk and current_chunk:
            chunks.append(current_chunk)
            current_chunk = []
            current_tokens = 0

        # If a single item is too large, split its content
        if item_tokens > max_tokens_per_chunk:
            # Split the content into smaller pieces
            content = item["content"]
            content_chunks = []
            chunk_size = len(content) * max_tokens_per_chunk // item_tokens

            for i in range(0, len(content), chunk_size):
                content_chunks.append(
                    {
                        "url": f"{item['url']} (part {i // chunk_size + 1})",
                        "content": content[i : i + chunk_size],
                    }
                )

            # Add the first part to current chunk if there's room
            if content_chunks:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = []
                    current_tokens = 0

                # Add each part as its own chunk
                for chunk_part in content_chunks:
                    chunks.append([chunk_part])
        else:
            current_chunk.append(item)
            current_tokens += item_tokens

    # Add the last chunk if it has content
    if current_chunk:
        chunks.append(current_chunk)

    return chunks
